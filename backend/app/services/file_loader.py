import os
import logging
from typing import List, Optional

from langchain.schema import Document

from app.services.ocr import run_ocr_on_image

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
}


def get_file_type(filename: str) -> Optional[str]:
    ext = os.path.splitext(filename.lower())[1]
    return SUPPORTED_EXTENSIONS.get(ext)


def is_supported_file(filename: str) -> bool:
    return get_file_type(filename) is not None


def _load_pdf(file_path: str) -> List[Document]:
    from app.services.pdf_loader import load_pdf
    return load_pdf(file_path)


def _load_docx(file_path: str) -> List[Document]:
    try:
        import docx as python_docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install with: pip install python-docx"
        )

    logger.info(f"[FileLoader] Loading DOCX: {file_path}")

    doc_obj = python_docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc_obj.paragraphs if p.text.strip()]

    if not paragraphs:
        raise ValueError(f"DOCX file '{os.path.basename(file_path)}' appears to be empty.")

    full_text = "\n\n".join(paragraphs)

    doc = Document(
        page_content=full_text,
        metadata={
            "source": file_path,
            "page": 0,
            "file_type": "docx",
        }
    )
    logger.info(f"[FileLoader] DOCX loaded: {len(paragraphs)} paragraph(s)")
    return [doc]


def _load_txt(file_path: str) -> List[Document]:
    logger.info(f"[FileLoader] Loading TXT: {file_path}")

    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError(
            f"Could not decode '{os.path.basename(file_path)}' "
            "— file may be binary, not plain text."
        )

    if not text:
        raise ValueError(f"Text file '{os.path.basename(file_path)}' is empty.")

    doc = Document(
        page_content=text,
        metadata={
            "source": file_path,
            "page": 0,
            "file_type": "txt",
        }
    )
    logger.info(f"[FileLoader] TXT loaded: {len(text)} chars")
    return [doc]


def _load_image(file_path: str) -> List[Document]:
    logger.info(f"[FileLoader] Loading image via OCR: {file_path}")
    return run_ocr_on_image(file_path)


def load_document(file_path: str) -> List[Document]:
    filename = os.path.basename(file_path)
    file_type = get_file_type(filename)

    if file_type is None:
        ext = os.path.splitext(filename)[1] or "(no extension)"
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    dispatch = {
        "pdf": _load_pdf,
        "docx": _load_docx,
        "txt": _load_txt,
        "image": _load_image,
    }

    loader_fn = dispatch[file_type]
    documents = loader_fn(file_path)

    if not documents:
        raise ValueError(f"No content could be extracted from '{filename}'.")

    logger.info(
        f"[FileLoader] '{filename}' → {len(documents)} document(s) "
        f"[type={file_type}]"
    )
    return documents